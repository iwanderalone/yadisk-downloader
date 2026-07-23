import os
import re
import logging
from datetime import datetime, timedelta
from io import BytesIO
import pandas as pd
import requests
from dotenv import load_dotenv

from telegram import ReplyKeyboardMarkup, Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    ConversationHandler,
    filters,
    ContextTypes,
)
from docx import Document
from docx.shared import Inches, Pt

from yandex_shared_disk import YandexSharedDiskClient, YandexSharedDiskError

# ==================== НАСТРОЙКИ ЛОГИРОВАНИЯ ====================
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# Отключаем отвлекающие системные логгер-запросы от библиотеки httpx / telegram (Long Polling)
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)
logging.getLogger("telegram").setLevel(logging.WARNING)

logger = logging.getLogger("TelegramYandexBot")

load_dotenv(override=True)

BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
YANDEX_OAUTH_TOKEN = os.getenv("YANDEX_OAUTH_TOKEN")
YANDEX_VD_HASH = os.getenv("YANDEX_VD_HASH")
REMOTE_FILE_PATH = os.getenv("YANDEX_FILE_PATH", "table.xlsx")
ADMIN_TELEGRAM_ID = os.getenv("ADMIN_TELEGRAM_ID", "791647629")

CACHE_LIFETIME = 300
WAIT_DATE_RANGE, WAIT_FEEDBACK = range(2)

_cache = {"df": None, "last_fetch": datetime.min}

# ==================== СЛОВАРЬ МУЛЬТИЯЗЫЧНОСТИ (RU / EN) ====================
TEXTS = {
    'RU': {
        'welcome_lang': "🌐 Пожалуйста, выберите язык / Please select your language:",
        'welcome_period': "👋 Привет! Выберите период:",
        'btn_today': "📅 Сегодня",
        'btn_week': "📆 За неделю",
        'btn_custom': "📅 Выбрать период",
        'btn_feedback': "✍️ Обратная связь",
        'btn_lang': "🌐 Сменить язык",
        'ask_range': "📅 Введите дату или диапазон дат, например:\n• 21.07.2026 - 23.07.2026\n• 21.07 - 23.07\n• 21.07",
        'ask_feedback': "✍️ Пожалуйста, напишите ваш отзыв, пожелание или сообщение об ошибке:",
        'feedback_thanks': "✅ Спасибо! Ваша обратная связь успешно отправлена.",
        'err_date_fmt': "❌ Не удалось распознать формат дат. Пожалуйста, введите в формате ДД.ММ.ГГГГ или ДД.ММ:",
        'cancel': "❌ Операция отменена.",
        'no_records': "📭 На выбранный период записей не найдено.",
        'caption': "📄 События за {}",
        'err_download': "❌ Ошибка скачивания с Яндекс Диска: {}",
        'label_today': "сегодня",
        'label_week': "неделю",
    },
    'EN': {
        'welcome_lang': "🌐 Please select your language / Пожалуйста, выберите язык:",
        'welcome_period': "👋 Hello! Please select a period:",
        'btn_today': "📅 Today",
        'btn_week': "📆 For the week",
        'btn_custom': "📅 Select period",
        'btn_feedback': "✍️ Send Feedback",
        'btn_lang': "🌐 Change language",
        'ask_range': "📅 Enter a date or a date range, for example:\n• 21.07.2026 - 23.07.2026\n• 21.07 - 23.07\n• 21.07",
        'ask_feedback': "✍️ Please write your feedback, suggestion, or bug report:",
        'feedback_thanks': "✅ Thank you! Your feedback has been successfully sent.",
        'err_date_fmt': "❌ Failed to parse date format. Please enter as DD.MM.YYYY or DD.MM:",
        'cancel': "❌ Operation cancelled.",
        'no_records': "📭 No records found for the selected period.",
        'caption': "📄 Events for {}",
        'err_download': "❌ Error downloading from Yandex Disk: {}",
        'label_today': "today",
        'label_week': "this week",
    }
}

def get_user_lang(context: ContextTypes.DEFAULT_TYPE) -> str:
    return context.user_data.get('lang', 'RU')

def get_period_keyboard(lang: str) -> ReplyKeyboardMarkup:
    t = TEXTS[lang]
    keyboard = [
        [t['btn_today'], t['btn_week']],
        [t['btn_custom'], t['btn_feedback']],
        [t['btn_lang']]
    ]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

def get_yandex_client() -> YandexSharedDiskClient:
    token = os.getenv("YANDEX_OAUTH_TOKEN") or YANDEX_OAUTH_TOKEN
    vd_hash = os.getenv("YANDEX_VD_HASH") or YANDEX_VD_HASH
    if not token or not vd_hash:
        raise YandexSharedDiskError(
            "Не заданы YANDEX_OAUTH_TOKEN или YANDEX_VD_HASH в файле .env!"
        )
    return YandexSharedDiskClient(oauth_token=token, vd_hash=vd_hash)

# ==================== УНИВЕРСАЛЬНЫЙ ПАРСЕР ДАТЫ ====================
def parse_date(val, sheet_name: str) -> datetime:
    if pd.isna(val):
        return None
    if isinstance(val, (datetime, pd.Timestamp)):
        return val
        
    val_str = str(val).strip()
    if not val_str:
        return None
        
    for fmt in ('%d.%m.%Y', '%d.%m.%y', '%Y-%m-%d'):
        try:
            return pd.to_datetime(val_str, format=fmt)
        except Exception:
            continue
            
    try:
        dt = pd.to_datetime(val_str, errors='raise')
        if dt.year <= 1900 or dt.year == 2001:
            inferred_year = datetime.now().year
            for word in sheet_name.split():
                if word.isdigit() and len(word) == 4:
                    inferred_year = int(word)
                    break
            dt = dt.replace(year=inferred_year)
        return dt
    except Exception:
        pass
        
    months_ru = {
        'янв': 1, 'фев': 2, 'мар': 3, 'апр': 4, 'май': 5, 'мая': 5,
        'июн': 6, 'июл': 7, 'авг': 8, 'сен': 9, 'окт': 10, 'ноя': 11, 'дек': 12
    }
    months_en = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5,
        'jun': 6, 'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    try:
        parts = val_str.lower().split()
        if len(parts) >= 2:
            day = int(parts[0])
            month_name = parts[1]
            month = None
            
            for k, v in months_ru.items():
                if month_name.startswith(k):
                    month = v
                    break
            if month is None:
                for k, v in months_en.items():
                    if month_name.startswith(k):
                        month = v
                        break
                        
            if month is not None:
                year = datetime.now().year
                if len(parts) >= 3 and parts[2].isdigit():
                    year = int(parts[2])
                else:
                    for word in sheet_name.split():
                        if word.isdigit() and len(word) == 4:
                            year = int(word)
                            break
                return pd.Timestamp(year=year, month=month, day=day)
    except Exception:
        pass
        
    return None

# ==================== ПАРСЕР ДИАПАЗОНА ДАТ ====================
def parse_user_date_range(text: str) -> tuple[datetime, datetime] | None:
    text = text.strip()
    date_pattern = r'(\d{1,2})[\.\-/](\d{1,2})(?:[\.\-/](\d{2,4}))?'
    matches = re.findall(date_pattern, text)
    
    if not matches:
        return None
        
    dates = []
    current_year = datetime.now().year
    
    for match in matches:
        day = int(match[0])
        month = int(match[1])
        year = int(match[2]) if match[2] else current_year
        if year < 100:
            year += 2000
        try:
            dates.append(datetime(year, month, day))
        except ValueError:
            return None
            
    if len(dates) == 1:
        return dates[0], dates[0]
    elif len(dates) >= 2:
        return min(dates[0], dates[1]), max(dates[0], dates[1])
        
    return None

# ==================== СКАЧИВАНИЕ С ОБЩЕГО ДИСКА ====================
def download_table() -> pd.DataFrame:
    now = datetime.now()
    if _cache["df"] is not None and (now - _cache["last_fetch"]).total_seconds() < CACHE_LIFETIME:
        logger.info("Использование кэшированных данных таблицы")
        return _cache["df"]

    client = get_yandex_client()
    file_path = os.getenv("YANDEX_FILE_PATH", REMOTE_FILE_PATH)
    logger.info("Запрос свежего файла '%s' с общего Яндекс Диска...", file_path)
    
    download_url = client.get_download_url(file_path)
    download_headers = {"Authorization": f"OAuth {client.oauth_token}"}
    resp = requests.get(download_url, headers=download_headers, timeout=60)
    resp.raise_for_status()

    if not resp.content.startswith(b'PK'):
        raise ValueError("Скачанный файл не является корректным XLSX.")

    try:
        xls = pd.ExcelFile(BytesIO(resp.content), engine='openpyxl')
        logger.info("Движок Excel: openpyxl")
    except Exception as e:
        logger.warning(f"openpyxl не справился: {e}. Пробую calamine...")
        try:
            xls = pd.ExcelFile(BytesIO(resp.content), engine='calamine')
            logger.info("Движок Excel: calamine")
        except Exception as e2:
            raise ValueError(f"Не удалось прочитать XLSX: {e2}")

    dfs = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        if 'Date' not in df.columns:
            continue
        df['Date'] = df['Date'].ffill()
        df.dropna(subset=['Date'], inplace=True)
        
        df['Date'] = df['Date'].apply(lambda val: parse_date(val, sheet_name))
        df.dropna(subset=['Date'], inplace=True)
        
        dfs.append(df)

    full_df = pd.concat(dfs, ignore_index=True) if dfs else pd.DataFrame()
    _cache["df"] = full_df
    _cache["last_fetch"] = now
    return full_df

# ==================== СОЗДАНИЕ WORD-ОТЧЕТА ====================
def create_docx(df: pd.DataFrame, label: str, lang: str = 'RU') -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    if df.empty:
        empty_msg = 'No records found for selected period.' if lang == 'EN' else 'Нет записей за выбранный период.'
        doc.add_paragraph(empty_msg)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    df_sorted = df.sort_values('Date')
    for date, group in df_sorted.groupby('Date'):
        date_para = doc.add_paragraph()
        run = date_para.add_run(date.strftime('%d.%m.%Y'))
        run.bold = True
        run.font.size = Pt(11)

        for _, row in group.iterrows():
            parts = []
            country = row.get('Country', '')
            event = row.get('Event', '')
            pkg = row.get('Original pkg', '')

            main = ''
            if pd.notna(country) and str(country).strip():
                main += str(country).strip()
            if pd.notna(event) and str(event).strip():
                main += f" — {str(event).strip()}" if main else str(event).strip()
            if pd.notna(pkg) and str(pkg).strip():
                main += f" ({str(pkg).strip()})" if main else str(pkg).strip()
            if not main:
                main = '[no title]' if lang == 'EN' else '[без названия]'
            parts.append(main)

            for col in df.columns:
                if col in ['Date', 'Country', 'Event', 'Original pkg']:
                    continue
                val = row[col]
                if pd.notna(val) and str(val).strip():
                    parts.append(f"{col}: {str(val).strip()}")
            event_text = "; ".join(parts)

            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(event_text)
            run.font.size = Pt(10)

        doc.add_paragraph()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==================== ОБРАБОТЧИКИ TELEGRAM ====================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """При старте бот сразу предлагает выбрать язык"""
    keyboard = [['🇷🇺 Русский', '🇬🇧 English']]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text(TEXTS['RU']['welcome_lang'], reply_markup=reply_markup)

async def set_language(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    if 'English' in text:
        lang = 'EN'
    else:
        lang = 'RU'
    
    context.user_data['lang'] = lang
    t = TEXTS[lang]
    await update.message.reply_text(t['welcome_period'], reply_markup=get_period_keyboard(lang))

async def handle_today(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return
    filtered = df[df['Date'] == today]
    await send_report(update, context, filtered, t['label_today'])

async def handle_week(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    start_date = today - timedelta(days=6)
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return
    mask = (df['Date'] >= start_date) & (df['Date'] <= today)
    await send_report(update, context, df[mask], t['label_week'])

async def choose_period_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    await update.message.reply_text(t['ask_range'])
    return WAIT_DATE_RANGE

async def receive_date_range(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    text = update.message.text.strip()
    
    parsed = parse_user_date_range(text)
    if not parsed:
        await update.message.reply_text(t['err_date_fmt'])
        return WAIT_DATE_RANGE
        
    start_date, end_date = parsed
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return ConversationHandler.END
        
    mask = (df['Date'] >= start_date) & (df['Date'] <= end_date)
    
    if start_date == end_date:
        label = start_date.strftime('%d.%m.%Y')
    else:
        label = f"{start_date.strftime('%d.%m.%Y')} – {end_date.strftime('%d.%m.%Y')}"
        
    await send_report(update, context, df[mask], label)
    return ConversationHandler.END

# ==================== ОБРАБОТЧИКИ ОБРАТНОЙ СВЯЗИ ====================
async def ask_feedback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.info(">>> [ЗАПРОС НА ОБРАТНУЮ СВЯЗЬ]")
    lang = get_user_lang(context)
    t = TEXTS[lang]
    await update.message.reply_text(t['ask_feedback'])
    return WAIT_FEEDBACK

async def receive_feedback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    feedback_text = update.message.text.strip()
    user = update.effective_user
    
    logger.info(">>> [ПОЛУЧЕН ОТЗЫВ]: '%s' от user_id=%s", feedback_text, user.id if user else "None")
    
    admin_id = os.getenv("ADMIN_TELEGRAM_ID", ADMIN_TELEGRAM_ID)
    if admin_id:
        username_str = f"@{user.username}" if user.username else "нет username"
        first_name = user.first_name or ""
        last_name = user.last_name or ""
        name_str = f"{first_name} {last_name}".strip() or "Без имени"
        
        admin_message = (
            f"🔔 *Получен новый отзыв!*\n\n"
            f"👤 *Отправитель:* {name_str} ({username_str})\n"
            f"🆔 *ID:* `{user.id}`\n\n"
            f"📝 *Текст отзыва:*\n{feedback_text}"
        )
        try:
            await context.bot.send_message(
                chat_id=admin_id,
                text=admin_message,
                parse_mode="Markdown"
            )
            logger.info("Отзыв успешно переслан администратору ID=%s", admin_id)
        except Exception as e:
            logger.error("Не удалось отправить отзыв администратору: %s", e)
            
    await update.message.reply_text(t['feedback_thanks'])
    return ConversationHandler.END

# ==================== ПРЯМОЙ ВВОД ДАТЫ (ВНЕ ДИАЛОГА) ====================
async def handle_direct_date_range(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    parsed = parse_user_date_range(text)
    if not parsed:
        return
        
    logger.info(">>> [ПРЯМОЙ ВВОД ДИАПАЗОНА ДАТ]: %s", text)
    lang = get_user_lang(context)
    t = TEXTS[lang]
    
    start_date, end_date = parsed
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return
        
    mask = (df['Date'] >= start_date) & (df['Date'] <= end_date)
    
    if start_date == end_date:
        label = start_date.strftime('%d.%m.%Y')
    else:
        label = f"{start_date.strftime('%d.%m.%Y')} – {end_date.strftime('%d.%m.%Y')}"
        
    await send_report(update, context, df[mask], label)

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    await update.message.reply_text(t['cancel'])
    return ConversationHandler.END

async def send_report(update: Update, context: ContextTypes.DEFAULT_TYPE, df: pd.DataFrame, label: str):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    if df.empty:
        await update.message.reply_text(t['no_records'])
        return
    bio = create_docx(df, label, lang=lang)
    bio.name = f"report_{label.replace(' ', '_')}.docx"
    await update.message.reply_document(document=bio, caption=t['caption'].format(label))

# ==================== ЗАПУСК ====================
def main():
    token = os.getenv("TELEGRAM_BOT_TOKEN") or BOT_TOKEN
    if not token:
        logger.error("TELEGRAM_BOT_TOKEN не задан! Пожалуйста, укажите его в .env файле.")
        return
    
    app = Application.builder().token(token).build()
    
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("lang", start))
    
    # Кнопки смены языка
    app.add_handler(MessageHandler(filters.Regex('^(🇷🇺 Русский|🇬🇧 English)$'), set_language))
    app.add_handler(MessageHandler(filters.Regex('^(🌐 Сменить язык|🌐 Change language)$'), start))

    # Кнопки выбора периода (RU + EN)
    app.add_handler(MessageHandler(filters.Regex('^(📅 Сегодня|📅 Today)$'), handle_today))
    app.add_handler(MessageHandler(filters.Regex('^(📆 За неделю|📆 For the week)$'), handle_week))

    # Обработчик прямого ввода даты (вне диалога)
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_direct_date_range))

    # Единый диалоговый менеджер для кастомного периода и фидбека
    conv_handler = ConversationHandler(
        entry_points=[
            MessageHandler(filters.Regex('^(📅 Выбрать период|📅 Select period)$'), choose_period_start),
            MessageHandler(filters.Regex('^(✍️ Обратная связь|✍️ Send Feedback)$'), ask_feedback)
        ],
        states={
            WAIT_DATE_RANGE: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_date_range)],
            WAIT_FEEDBACK: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_feedback)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    app.add_handler(conv_handler)

    logger.info("Telegram бот успешно запущен и готов к работе (RU / EN)")
    app.run_polling()

if __name__ == "__main__":
    main()
