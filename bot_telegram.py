import os
import logging
from datetime import datetime, timedelta
from io import BytesIO
import pandas as pd
import requests
from dotenv import load_dotenv

from telegram import ReplyKeyboardMarkup, Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    ConversationHandler,
    filters,
    ContextTypes,
)
from docx import Document
from docx.shared import Inches, Pt

from yandex_shared_disk import YandexSharedDiskClient, YandexSharedDiskError

# ==================== НАСТРОЙКИ ЛОГИРОВАНИЯ ====================
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# Отключаем отвлекающие системные логгер-запросы от библиотеки httpx / telegram (Long Polling)
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)
logging.getLogger("telegram").setLevel(logging.WARNING)

logger = logging.getLogger("TelegramYandexBot")

# Загружаем .env с принудительным переопределением (override=True)
load_dotenv(override=True)

BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
YANDEX_OAUTH_TOKEN = os.getenv("YANDEX_OAUTH_TOKEN")
YANDEX_VD_HASH = os.getenv("YANDEX_VD_HASH")
REMOTE_FILE_PATH = os.getenv("YANDEX_FILE_PATH", "table.xlsx")

CACHE_LIFETIME = 300
WAIT_START_DATE, WAIT_END_DATE = range(2)

_cache = {"df": None, "last_fetch": datetime.min}

# ==================== СЛОВАРЬ МУЛЬТИЯЗЫЧНОСТИ (RU / EN) ====================
TEXTS = {
    'RU': {
        'welcome_lang': "🌐 Пожалуйста, выберите язык / Please select your language:",
        'welcome_period': "👋 Привет! Выберите период:",
        'btn_today': "📅 Сегодня",
        'btn_week': "📆 За неделю",
        'btn_custom': "📅 Выбрать период",
        'btn_lang': "🌐 Сменить язык",
        'ask_start': "📅 Введите начальную дату (ДД.ММ.ГГГГ):",
        'ask_end': "📅 Введите конечную дату (ДД.ММ.ГГГГ):",
        'err_date_fmt': "❌ Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ:",
        'err_date_order': "❌ Конечная дата раньше начальной. Введите еще раз:",
        'cancel': "❌ Выбор периода отменён.",
        'no_records': "📭 На выбранный период записей не найдено.",
        'caption': "📄 События за {}",
        'err_download': "❌ Ошибка скачивания с Яндекс Диска: {}",
        'label_today': "сегодня",
        'label_week': "неделю",
    },
    'EN': {
        'welcome_lang': "🌐 Please select your language / Пожалуйста, выберите язык:",
        'welcome_period': "👋 Hello! Please select a period:",
        'btn_today': "📅 Today",
        'btn_week': "📆 For the week",
        'btn_custom': "📅 Select period",
        'btn_lang': "🌐 Change language",
        'ask_start': "📅 Enter start date (DD.MM.YYYY):",
        'ask_end': "📅 Enter end date (DD.MM.YYYY):",
        'err_date_fmt': "❌ Invalid date format. Enter as DD.MM.YYYY:",
        'err_date_order': "❌ End date is before start date. Enter again:",
        'cancel': "❌ Period selection cancelled.",
        'no_records': "📭 No records found for the selected period.",
        'caption': "📄 Events for {}",
        'err_download': "❌ Error downloading from Yandex Disk: {}",
        'label_today': "today",
        'label_week': "this week",
    }
}

def get_user_lang(context: ContextTypes.DEFAULT_TYPE) -> str:
    return context.user_data.get('lang', 'RU')

def get_period_keyboard(lang: str) -> ReplyKeyboardMarkup:
    t = TEXTS[lang]
    keyboard = [
        [t['btn_today'], t['btn_week']],
        [t['btn_custom'], t['btn_lang']]
    ]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

def get_yandex_client() -> YandexSharedDiskClient:
    token = os.getenv("YANDEX_OAUTH_TOKEN") or YANDEX_OAUTH_TOKEN
    vd_hash = os.getenv("YANDEX_VD_HASH") or YANDEX_VD_HASH
    if not token or not vd_hash:
        raise YandexSharedDiskError(
            "Не заданы YANDEX_OAUTH_TOKEN или YANDEX_VD_HASH в файле .env!"
        )
    return YandexSharedDiskClient(oauth_token=token, vd_hash=vd_hash)

# ==================== СКАЧИВАНИЕ С ОБЩЕГО ДИСКА ====================
def download_table() -> pd.DataFrame:
    now = datetime.now()
    if _cache["df"] is not None and (now - _cache["last_fetch"]).total_seconds() < CACHE_LIFETIME:
        logger.info("Использование кэшированных данных таблицы")
        return _cache["df"]

    client = get_yandex_client()
    file_path = os.getenv("YANDEX_FILE_PATH", REMOTE_FILE_PATH)
    logger.info("Запрос свежего файла '%s' с общего Яндекс Диска...", file_path)
    
    download_url = client.get_download_url(file_path)
    download_headers = {"Authorization": f"OAuth {client.oauth_token}"}
    resp = requests.get(download_url, headers=download_headers, timeout=60)
    resp.raise_for_status()

    if not resp.content.startswith(b'PK'):
        raise ValueError("Скачанный файл не является корректным XLSX.")

    try:
        xls = pd.ExcelFile(BytesIO(resp.content), engine='openpyxl')
        logger.info("Движок Excel: openpyxl")
    except Exception as e:
        logger.warning(f"openpyxl не справился: {e}. Пробую calamine...")
        try:
            xls = pd.ExcelFile(BytesIO(resp.content), engine='calamine')
            logger.info("Движок Excel: calamine")
        except Exception as e2:
            raise ValueError(f"Не удалось прочитать XLSX: {e2}")

    dfs = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        if 'Date' not in df.columns:
            continue
        df['Date'] = df['Date'].ffill()
        df.dropna(subset=['Date'], inplace=True)
        df['Date'] = pd.to_datetime(df['Date'], format='%d.%m.%Y', errors='coerce')
        df.dropna(subset=['Date'], inplace=True)
        dfs.append(df)

    full_df = pd.concat(dfs, ignore_index=True) if dfs else pd.DataFrame()
    _cache["df"] = full_df
    _cache["last_fetch"] = now
    return full_df

# ==================== СОЗДАНИЕ WORD-ОТЧЕТА ====================
def create_docx(df: pd.DataFrame, label: str, lang: str = 'RU') -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    if df.empty:
        empty_msg = 'No records found for selected period.' if lang == 'EN' else 'Нет записей за выбранный период.'
        doc.add_paragraph(empty_msg)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    df_sorted = df.sort_values('Date')
    for date, group in df_sorted.groupby('Date'):
        date_para = doc.add_paragraph()
        run = date_para.add_run(date.strftime('%d.%m.%Y'))
        run.bold = True
        run.font.size = Pt(11)

        for _, row in group.iterrows():
            parts = []
            country = row.get('Country', '')
            event = row.get('Event', '')
            pkg = row.get('Original pkg', '')

            main = ''
            if pd.notna(country) and str(country).strip():
                main += str(country).strip()
            if pd.notna(event) and str(event).strip():
                main += f" — {str(event).strip()}" if main else str(event).strip()
            if pd.notna(pkg) and str(pkg).strip():
                main += f" ({str(pkg).strip()})" if main else str(pkg).strip()
            if not main:
                main = '[no title]' if lang == 'EN' else '[без названия]'
            parts.append(main)

            for col in df.columns:
                if col in ['Date', 'Country', 'Event', 'Original pkg']:
                    continue
                val = row[col]
                if pd.notna(val) and str(val).strip():
                    parts.append(f"{col}: {str(val).strip()}")
            event_text = "; ".join(parts)

            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(event_text)
            run.font.size = Pt(10)

        doc.add_paragraph()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==================== ОБРАБОТЧИКИ TELEGRAM ====================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    logger.info(">>> [ПОЛУЧЕНО СООБЩЕНИЕ /start] от user_id=%s (@%s)", user.id if user else "None", user.username if user else "None")
    keyboard = [['🇷🇺 Русский', '🇬🇧 English']]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text(TEXTS['RU']['welcome_lang'], reply_markup=reply_markup)

async def set_language(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    logger.info(">>> [ВЫБОР ЯЗЫКА]: %s", text)
    if 'English' in text:
        lang = 'EN'
    else:
        lang = 'RU'
    
    context.user_data['lang'] = lang
    t = TEXTS[lang]
    await update.message.reply_text(t['welcome_period'], reply_markup=get_period_keyboard(lang))

async def handle_today(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.info(">>> [ЗАПРОС СЕГОДНЯ]")
    lang = get_user_lang(context)
    t = TEXTS[lang]
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return
    filtered = df[df['Date'] == today]
    await send_report(update, context, filtered, t['label_today'])

async def handle_week(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.info(">>> [ЗАПРОС ЗА НЕДЕЛЮ]")
    lang = get_user_lang(context)
    t = TEXTS[lang]
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    start_date = today - timedelta(days=6)
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return
    mask = (df['Date'] >= start_date) & (df['Date'] <= today)
    await send_report(update, context, df[mask], t['label_week'])

async def choose_period_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.info(">>> [ЗАПРОС СВОЙ ПЕРИОД]")
    lang = get_user_lang(context)
    t = TEXTS[lang]
    await update.message.reply_text(t['ask_start'])
    return WAIT_START_DATE

async def receive_start_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    text = update.message.text.strip()
    try:
        start_date = datetime.strptime(text, '%d.%m.%Y')
    except ValueError:
        await update.message.reply_text(t['err_date_fmt'])
        return WAIT_START_DATE
    context.user_data['start_date'] = start_date
    await update.message.reply_text(t['ask_end'])
    return WAIT_END_DATE

async def receive_end_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    text = update.message.text.strip()
    try:
        end_date = datetime.strptime(text, '%d.%m.%Y')
    except ValueError:
        await update.message.reply_text(t['err_date_fmt'])
        return WAIT_END_DATE
    start_date = context.user_data['start_date']
    if end_date < start_date:
        await update.message.reply_text(t['err_date_order'])
        return WAIT_END_DATE
    try:
        df = download_table()
    except Exception as e:
        await update.message.reply_text(t['err_download'].format(e))
        return ConversationHandler.END
    mask = (df['Date'] >= start_date) & (df['Date'] <= end_date)
    label = f"{start_date.strftime('%d.%m.%Y')} – {end_date.strftime('%d.%m.%Y')}"
    await send_report(update, context, df[mask], label)
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    await update.message.reply_text(t['cancel'])
    return ConversationHandler.END

async def send_report(update: Update, context: ContextTypes.DEFAULT_TYPE, df: pd.DataFrame, label: str):
    lang = get_user_lang(context)
    t = TEXTS[lang]
    if df.empty:
        await update.message.reply_text(t['no_records'])
        return
    bio = create_docx(df, label, lang=lang)
    bio.name = f"report_{label.replace(' ', '_')}.docx"
    await update.message.reply_document(document=bio, caption=t['caption'].format(label))

async def any_message_log(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text if update.message else "None"
    user = update.effective_user
    logger.info(">>> [ЛЮБОЕ ВХОДЯЩЕЕ СООБЩЕНИЕ]: '%s' от user_id=%s (@%s)", text, user.id if user else "None", user.username if user else "None")

# ==================== ЗАПУСК ====================
def main():
    token = os.getenv("TELEGRAM_BOT_TOKEN") or BOT_TOKEN
    if not token:
        logger.error("TELEGRAM_BOT_TOKEN не задан! Пожалуйста, укажите его в .env файле.")
        return
        
    logger.info("Запуск Telegram бота...")
    app = Application.builder().token(token).build()
    
    # Логгируем абсолютно все входящие текстовые сообщения на верхнем уровне
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, any_message_log), group=-1)

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("lang", start))
    
    app.add_handler(MessageHandler(filters.Regex('^(🇷🇺 Русский|🇬🇧 English)$'), set_language))
    app.add_handler(MessageHandler(filters.Regex('^(🌐 Сменить язык|🌐 Change language)$'), start))

    app.add_handler(MessageHandler(filters.Regex('^(📅 Сегодня|📅 Today)$'), handle_today))
    app.add_handler(MessageHandler(filters.Regex('^(📆 За неделю|📆 For the week)$'), handle_week))

    conv_handler = ConversationHandler(
        entry_points=[MessageHandler(filters.Regex('^(📅 Выбрать период|📅 Select period)$'), choose_period_start)],
        states={
            WAIT_START_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_start_date)],
            WAIT_END_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_end_date)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    app.add_handler(conv_handler)

    logger.info("Telegram бот успешно запущен и готов к работе (RU / EN)")
    app.run_polling()

if __name__ == "__main__":
    main()
